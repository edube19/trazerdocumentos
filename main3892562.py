from pydoc import cli, doc
import sys
#from this import s
from xmlrpc import client
from docx import Document
from docx.shared import Cm
import docx
from recursos import*

if __name__ == "__main__":
    while True:
        cliente='LOURDES HIDALGO CASTILLO'
        clienteInmueble='LOURDES'
        empleador='CIMANDESPERU'
        conyuge1='DIAZ CASTRO AUREA FARINA'
        conyuge2='LUNA CORDOVA HENRY JOHN'

        conyuge1v2='AUREA FARINA'
        conyuge2v2='HENRY JOHN'

        try:
            #COPIA LITERAL, NO ESTA NINGUNO DE LOS NOMBRES
            ruta1='C:/Users/DELL/Desktop/pruebaword/exp3892562/copia literal/Copia literal de partida electronica del inmueble_3892562_1.pdf'
            #pag 4 y 5
            extraerTexto_imagenV2('Copia literal de partida electronica del inmueble_3892562_1_page-',6,conyuge1v2,'copia literal','exp3892562')

            #pag 3 y 4
            extraerTexto_imagenV2('Copia literal de partida electronica del inmueble_3892562_1_page-',6,conyuge2v2,'copia literal','exp3892562')
        except Exception as e:
            print('Error en COPIA LITERAL: ' + str(e))

        """try:
        #HR PU
            #se encontro en la pag 7
            extraerTexto_imagenV2('HR y PU o constancia de no Adeudo por impuesto predial original del inmueble_3892562_1_page-',7,conyuge1,'hr pu','exp3892562')
            #no se encontro
            extraerTexto_imagenV2('HR y PU o constancia de no Adeudo por impuesto predial original del inmueble_3892562_1_page-',7,conyuge2,'hr pu','exp3892562')
            
            #se encontro en la pag 3,5,6,7
            extraerTexto_imagenV2('INFTASCIVALID_3892626_1_TASACIONANTICIPADA_22052021_page-',7,conyuge1,'hr pu','exp3892562')
            #no se encontro
            extraerTexto_imagenV2('INFTASCIVALID_3892626_1_TASACIONANTICIPADA_22052021_page-',7,conyuge2,'hr pu','exp3892562')
            
        except Exception as e:
            print('Error en HR PU: '+str(e))"""

        """try:
        #MINUTA
        #ruta2='C:/Users/DELL/Desktop/pruebaword/exp3892562/minuta/Minuta de Compra Venta_3892562_5.docx'
        #doc2=docx.Document(ruta2)
        #buscar_palabra(cliente,doc2)

        #ruta3='C:/Users/DELL/Desktop/pruebaword/exp3892562/minuta/SOLCREDITOHIPOTECARIO_3892562_6.docx'
        #doc3=docx.Document(ruta3)
        #buscar_palabra(cliente,doc3)
            pass
        except Exception as e:
            print('Error en MINUTA: '+str(e))

        try:
        #PRESTAMOFIRMA

        #ruta4='C:/Users/DELL/Desktop/pruebaword/exp3892562/prestamofirma/SOLCREDITOHIPOTECARIO_3892562_15.docx'
        #doc4=docx.Document(ruta4)
        #buscar_palabra(cliente,doc4)
            pass
        except Exception as e:
            print('Error en PRESTAMOFIRMA: '+str(e))

        try:
        #TITULO
            ruta5='C:/Users/DELL/Desktop/pruebaword/exp3892562/titulo/SOLCREDITOHIPOTECARIO_3892562_11.docx'
            doc5=docx.Document(ruta5)
            buscar_palabra(conyuge1v2,doc5)
            buscar_palabra(conyuge2v2,doc5)
        except Exception as e:
            print('Error en TITULO: '+str(e))"""
            
        sys.exit()

