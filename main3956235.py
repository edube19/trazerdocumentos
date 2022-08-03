from pydoc import doc
import sys
from docx import Document
from docx.shared import Cm
import docx
from recursos import*

if __name__ == "__main__":
    while True:
        
        cliente='RENZO ZEBALLOS DEZA'
        clienteInmueble='RENZO'
        empleador='CORPORACION SECURITY TECH'

        """try:
        #TITULO
            ruta1='C:/Users/DELL/Desktop/pruebaword/exp3956235/titulo/SOLCREDITOHIPOTECARIO_3956235_14.docx'
            doc1 = docx.Document(ruta1)
            buscar_palabra(cliente,doc1)
        except Exception as e:
            print("Error en TITULO: " + str(e))

        try:
        #EVIDENCIA ABONO
            ruta2='C:/Users/DELL/Desktop/pruebaword/exp3956235/evidencia de abono/Evidencia de abonos en cuenta_3956235_1.pdf'
            leerpdf(ruta2,empleador)
        except Exception as e:
            print("Error en EVIDENCIA ABONO: " + str(e))"""

        try:
        #GARANTIA mobiliaria
            ruta3='C:/Users/DELL/Desktop/pruebaword/exp3956235/garantia mobiliaria/Licencia de funcionamiento o contrato de alquiler o HR PU del local en el que trabaja_3956235_1.pdf'
            leerpdf(ruta3,empleador)

            ruta4='C:/Users/DELL/Desktop/pruebaword/exp3956235/garantia mobiliaria/SOLCREDITOHIPOTECARIO_3956235_13.docx'
            doc4 = docx.Document(ruta4)
            buscar_palabra(cliente,doc4)
        except Exception as e:
            print('Error en GARANTIA mobiliaria:'+str(e))

        try:    
        #LEVANTAMIENTO Hipotecaria
            ruta5='C:/Users/DELL/Desktop/pruebaword/exp3956235/hipoteca/Declaracion Jurada Hipotecaria_3956235_1.pdf'
            leerpdf(ruta5,cliente)
            
            ruta6='C:/Users/DELL/Desktop/pruebaword/exp3956235/hipoteca/OtrosCI2_3956235_2.pdf'
            leerpdf(ruta6,cliente)
            
            ruta7='C:/Users/DELL/Desktop/pruebaword/exp3956235/hipoteca/SOLCREDITOHIPOTECARIO_3956235_14.docx'
            doc7 = docx.Document(ruta7)
            buscar_palabra(cliente,doc7)

            ruta8='C:/Users/DELL/Desktop/pruebaword/exp3956235/hipoteca/Solicitud de Credito Hipotecario_3956235_3.pdf'
            leerpdf(ruta8,clienteInmueble)

            ruta9='C:/Users/DELL/Desktop/pruebaword/exp3956235/hipoteca/Solicitud de Seguro de Desgravamen_3956235_1.pdf'
            leerpdf(ruta9,clienteInmueble)

            ruta10='C:/Users/DELL/Desktop/pruebaword/exp3956235/hipoteca/Solicitud de seguro de incendio de uso vivienda o uso mixto_3956235_1.pdf'
            leerpdf(ruta10,clienteInmueble)

            ruta11='C:/Users/DELL/Desktop/pruebaword/exp3956235/minuta/Minuta de Compra Venta_3956235_2 (2)_pages-to-jpg-1.jpg'
            extraerTexto_imagen(cliente,ruta11)
        except Exception as e:
            print('Error en LEVANTAMIENTO Hipotecaria:'+str(e))
        """
        try:
        #MINUTA
            extraerTexto_imagenV2('Minuta de Compra Venta_3956235_1 (1)_page-',2,cliente,'minuta','exp3956235')
            extraerTexto_imagenV2('Minuta de Compra Venta_3956235_2 (2)_pages-to-jpg-',13,cliente,'minuta','exp3956235')#1 vez
            extraerTexto_imagenV2('Minuta de Compra Venta_3991131_1_TASACIONANTICIPADA_10062022 (1)_page-',13,cliente,'minuta','exp3956235')#1 vez
        except Exception as e:
            print('Error en MINUTA: '+ str(e))"""
        sys.exit()