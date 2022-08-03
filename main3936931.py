from pydoc import cli, doc
import sys
from xmlrpc import client
from docx import Document
from docx.shared import Cm
import docx
from recursos import*

if __name__ == "__main__":
    while True:
        cliente ='CLAUDETT KATERINA DELGADO LLANOS'
        clientev2='Claudett Katerina Delgado Llanos'
        empleador='JURADO NACIONAL DE ELECCIONES'
        clienteInmueble='CLAUDETT'
        constructora='PROMOTORA LOS ALAMOS'

        try:
        #ADENDA
            ruta1='C:/Users/DELL/Desktop/pruebaword/exp3936931/adenda/Minuta de Compra Venta_3936931_2.docx'
            doc1=docx.Document(ruta1)
            buscar_palabra(cliente, doc1)
            
            ruta2='C:/Users/DELL/Desktop/pruebaword/exp3936931/adenda/SOLCREDITOHIPOTECARIO_3936931_9.docx'
            doc2=docx.Document(ruta2)
            buscar_palabra(cliente, doc2)

            #se econtro en la pag2
            extraerTexto_imagenV2('Solicitud de Credito Hipotecario_3936931_17_page-',4,cliente,'adenda','exp3936931')
        except Exception as e:
            print('Error en ADENDA :'+str(e))

        try:    
        #GARANTIA MOBILIARIA
            ruta3='C:/Users/DELL/Desktop/pruebaword/exp3936931/garantiamobiliaria/SOLCREDITOHIPOTECARIO_3936931_20.docx'
            doc3=docx.Document(ruta3)
            buscar_palabra(cliente, doc3)
        except:
            print('Error en GARANTIA MOBILIARIA: '+str(e))

        try:
        #MINUTA
        #se econtro en la pag9
            extraerTexto_imagenV2('COPSIMPLEHR_3037813_1_TASACIONANTICIPADA_19062022_page-',19,clientev2,'minuta','exp3936931')

            ruta4='C:/Users/DELL/Desktop/pruebaword/exp3936931/minuta/Minuta de Compra Venta_3936931_2.docx'
            doc4=docx.Document(ruta4)
            buscar_palabra(cliente, doc4)

            #se encontro en la pag9 y pag10
            #extraerTexto_imagenV2('Minuta de Compra Venta_3996002_1_TASACIONANTICIPADA_27062022_page-',14,cliente,'minuta','exp3936931')

            leerpdf('C:/Users/DELL/Desktop/pruebaword/exp3936931/minuta/SOLCREDITOHIPOTECARIO_3936931_6.pdf',constructora)
        
            extraerTexto_imagenV2('Solicitud de Credito Hipotecario_3936931_7_page-',14,constructora,'minuta','exp3936931')
        except:
            print('Error en MINUTA: '+str(e))

        try:
        #PRESTAMOFIRMA
            ruta5='C:/Users/DELL/Desktop/pruebaword/exp3936931/prestamofirma/OtrosCI_3936931_1.docx'
            doc5=docx.Document(ruta5)
            buscar_palabra(cliente, doc5)

            ruta6='C:/Users/DELL/Desktop/pruebaword/exp3936931/prestamofirma/SOLCREDITOHIPOTECARIO_3936931_22.docx'
            doc6=docx.Document(ruta6)
            buscar_palabra(cliente, doc6)
        except:
            print('Error en PRESTAMO FIRMA: '+str(e))

        try:    
        #TITULO
            ruta7='C:/Users/DELL/Desktop/pruebaword/exp3936931/titulo/SOLCREDITOHIPOTECARIO_3936931_11.docx'
            doc7=docx.Document(ruta7)
            buscar_palabra(cliente, doc7)

            ruta8='C:/Users/DELL/Desktop/pruebaword/exp3936931/titulo/SOLCREDITOHIPOTECARIO_3936931_16.docx'
            doc8=docx.Document(ruta8)
            buscar_palabra(cliente, doc8)

            ruta9='C:/Users/DELL/Desktop/pruebaword/exp3936931/titulo/SOLCREDITOHIPOTECARIO_3936931_19.docx'
            doc9=docx.Document(ruta9)
            buscar_palabra(cliente, doc9)
        except:
            print('Error en TITULO: '+str(e))

        sys.exit()