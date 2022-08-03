from pydoc import cli, doc
import sys
from xmlrpc import client
from docx import Document
from docx.shared import Cm
import docx
from recursos import*

if __name__ == "__main__":
    while True:
        cliente = 'JULIO CESAR LABAN RAMIREZ'
        clientev2='JULIO'
        clientev3='LABAN RAMIREZ JULIO CESAR'
        empleador = 'J & V RESGUARDO'
        constructora= 'MIRANDA CONSTRUCTORES'
        dni='6046590'

        """try:
        #CLAUSULA ADICIONAL 
            leerpdf('C:/Users/DELL/Desktop/pruebaword/exp3974679/clausula adicional/Minuta de Compra Venta_3974679_3.pdf',clientev2)
            leerpdf('C:/Users/DELL/Desktop/pruebaword/exp3974679/clausula adicional/Minuta de Compra Venta_3993816_3_TASACIONANTICIPADA_17062022.pdf',clientev2)
        except Exception as e:
            print('Error en CLAUSULA ADICIONAL: '+str(e))

        try:    
        #EVIDENCIA
            extraerTexto_imagenV2('Copia de Convenio RIA_3974679_1_page-',2,cliente,'evidencia','exp3974679')
            extraerTexto_imagenV2('Evidencia Aprobacion Tasa_3974679_1_page-',1,dni,'evidencia','exp3974679')
            extraerTexto_imagenV2('SOLCREDITOHIPOTECARIO_3974679_2_page-',1,clientev3,'evidencia','exp3974679')
        except Exception as e:
            print('Error en EVIDENCIA: '+str(e))"""

        try:
        #GARANTIA MOBILIARIA
            ruta1='C:/Users/DELL/Desktop/pruebaword/exp3974679/garantiamob/SOLCREDITOHIPOTECARIO_3974679_10.docx'
            doc1=docx.Document(ruta1)
            buscar_palabra(cliente,doc1)
        except Exception as e:
            print('Error en GARANTIA MOBILIARIA: '+str(e))
        """
        try:
        #HIPOTECA
            extraerTexto_imagenV2('Declaracion Jurada Hipotecaria_3974679_1_page-',3,cliente,'hipoteca','exp3974679')
        except Exception as e:
            print('Error en HIPOTECA: '+str(e))

        try:
        #INMUEBLE
            leerpdf('C:/Users/DELL/Desktop/pruebaword/exp3974679/inmueble/Solicitud de seguro de incendio de uso vivienda o uso mixto_3974679_1.pdf',clientev2)
            leerpdf('C:/Users/DELL/Desktop/pruebaword/exp3974679/inmueble/SOLSEGUROINCENDIO_3974679_1.pdf',clientev2)
        except Exception as e:
            print('Error en INMUEBLE: '+str(e))

        try:
        #MINUTA
            leerpdf('C:/Users/DELL/Desktop/pruebaword/exp3974679/minuta/Minuta de Compra Venta_3974679_5.pdf',constructora)
            leerpdf('C:/Users/DELL/Desktop/pruebaword/exp3974679/minuta/Minuta de Compra Venta_3993816_3_TASACIONANTICIPADA_17062022.pdf',constructora)
        except Exception as e:
            print('Error en MINUTA: '+str(e))

        try:
        #PRESTAMO FIRMA
            ruta2='C:/Users/DELL/Desktop/pruebaword/exp3974679/prestamofirma/SOLCREDITOHIPOTECARIO_3974679_9.docx'
            doc2=docx.Document(ruta2)
            buscar_palabra(cliente,doc2)
        except Exception as e:
            print('Error en FIRMA: '+str(e))

        try:
        #TITULO
            ruta3='C:/Users/DELL/Desktop/pruebaword/exp3974679/titulo/SOLCREDITOHIPOTECARIO_3974679_12.docx'
            doc3=docx.Document(ruta3)
            buscar_palabra(cliente,doc3)
        except Exception as e:
            print('Error en TITULO: '+str(e))"""
        sys.exit()