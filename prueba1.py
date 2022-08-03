from calendar import c
from docx import Document
from docx.shared import Cm
import docx

ruta='SOLCREDITOHIPOTECARIO_3956235_13.docx'
doc = docx.Document(ruta)
l=len(doc.paragraphs)
print(l)
#son 530
buscar='RENZO ZEBALLOS DEZA'
for i in range(l):
    cadena=doc.paragraphs[i].text
    r=cadena.find(buscar)
    if r!=-1:
        print('Se encontro en la linea '+str(i+1))
        print(cadena)
    #print(doc.paragraphs[i].text)

"""  
lp=len(doc.paragraphs[1].runs)
print (lp)

for j in range(lp):
    print(doc.paragraphs[1].runs[j].text)"""  