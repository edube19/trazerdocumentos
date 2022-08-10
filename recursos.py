from array import array
from audioop import add
from calendar import c
from cgi import print_arguments
from re import T, sub
from turtle import st
from docx import Document
import aspose.words as aw
from fitz import *
from email.mime import image
from importlib.resources import path
from PIL import Image
from pyparsing import condition_as_parse_action 
from pytesseract import pytesseract 
from pdf2image import convert_from_path
from requests import delete
from recurso_prueba import*
import docx
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def principalv3(ruta,string,ruta_guardar):
    #para que almacene el documento en una variable
    doc=docx.Document(ruta)

    if (buscar_palabra(string,doc)):#buscar si el nombre del cliente se encuentra en el documento
        
        #U+200E es un caracter vacio, solo para "borrar" la palabra SEÑOR NOTARIO
        editar_linea(doc,'SEÑOR NOTARIO:','‎',ruta_guardar)

        #poner en literal la cantidad de porcentaje
        porcentaje_decimalv2(doc,ruta_guardar)

        #cambiar las comas por puntos
        cambiarcomas(doc,ruta_guardar)

        #en caso haya cantidades de mas de millon
        camrbiarmillon(doc,ruta_guardar)
    else:
        print('El cliente no corresponde al documento')

def reconstruir_string(string):
    string=string.strip()#elimina los espacios en blanco del comienzo y final
    ns = " ".join(string.split())
    return ns
    
def modificartamano(ruta,string,ruta_guardar):#en construccion
    doc=docx.Document(ruta)
    ruta_aux='C:/Users/DELL/Desktop/pruebaword/modificaciones.docx'
    l=cantidad_lineas(doc)
    a=-1
    c=0
    cadenafinal=[]
    doc=docx.Document(ruta)
    doc_aux=docx.Document(ruta_aux)
    for i in range(l):
        cadena=doc.paragraphs[i].text
        r=cadena.find(string) 
        if r!=a:
            guardar_cadena=cadena.split(string)
            para=doc_aux.add_paragraph(string)
            para.font.size = Pt(12)
            string_nuevo=para
            for j in range(0,len(guardar_cadena),2):
                subcadena=guardar_cadena[j]
                cadenafinal[j]=subcadena
                try:
                    cadenafinal[j+1]=string_nuevo
                except Exception as e:
                    pass
            c=c+1
        cadenafinal.clear()
    doc.save(ruta_guardar)
    if (c==0):
        print('No se encontro')

def leerdoc(doc):
    l=cantidad_lineas(doc)
    for i in range(l):
        cadena=doc.paragraphs[i].text
        print('Parrafo N '+str(i+1))
        print(cadena)

def eliminar_linea(doc,string,ruta_guardar):
    l=cantidad_lineas(doc)
    a=-1
    c=0
    for i in range(l):
        cadena=doc.paragraphs[i].text
        r=cadena.find(string)
        if r!=a:
            del doc.paragraphs[i]
            c=c+1
            doc.save(ruta_guardar)
            break
    if (c==0):
        print('No se encontro')

def porcentaje_decimalv2(doc,ruta_guardar):
    l=cantidad_lineas(doc)
    a=-1
    for i in range(l):
        cadena=doc.paragraphs[i].text
        if (evitar_sobreescritura(cadena,'por ciento')):
            r=cadena.find('%')
            if (r!=a):
                guardar_cadena=cadena.split('%')
                for j in range(len(guardar_cadena)-1):
                    numero=''
                    contador=1
                    cond=True
                    subcadena=guardar_cadena[j]
                    l=len(subcadena)
                    while(cond):
                        valor=subcadena[l-contador]#obtenner el valor
                        cond=valor.isdigit()#comprobar si es digito o no, si es falso acaba la ejecucion
                        if (cond):
                            numero=valor+numero#guardando el numero
                            contador=contador+1#avanza para el siguiente valor
                        elif(valor=='.'):# para casos → ab.cd%
                            contador=contador+1
                            numero='.'+numero
                            cond=True
                        elif (valor==' ' and numero==''):# para casos → ab.cd %
                            contador=contador+1
                            cond=True
                    try:
                        punto=numero.find('.')
                        if punto!=a:
                            numero_separado=numero.split('.')#separa el string segun el caracter separador
                            long_decimal=len(numero_separado[1])
                            cero='0' 
                            decimal_cero=cero*long_decimal
                            if (decimal_cero!=numero_separado[1]):
                                numeroconv1=int(numero_separado[0])#parte entera
                                numeroconv2=int(numero_separado[1])#parte decimal
                                porcentaje1=numero_to_letras(numeroconv1)
                                porcentaje2=numero_to_letras(numeroconv2)
                                porcentajeminus1=porcentaje1.lower()
                                porcentajeminus2=porcentaje2.lower()
                                nuevostring='% ('+porcentajeminus1+' punto '+porcentajeminus2+' por ciento)'
                                guardar_cadena[j]=guardar_cadena[j]+nuevostring
                            else:
                                numeroconv1=int(numero_separado[0])#parte entera
                                porcentaje1=numero_to_letras(numeroconv1)
                                porcentajeminus1=porcentaje1.lower()
                                nuevostring='% ('+porcentajeminus1+' por ciento)'
                                guardar_cadena[j]=guardar_cadena[j]+nuevostring
                            #doc.paragraphs[i].text=nueva_cadena
                        else:
                            numeroconv1=int(numero)
                            porcentaje1=numero_to_letras(numeroconv1)
                            porcentajeminus1=porcentaje1.lower()
                            nuevostring='% ('+porcentajeminus1+' por ciento)'
                            guardar_cadena[j]=guardar_cadena[j]+nuevostring
                            #doc.paragraphs[i].text=nueva_cadena
                    except Exception as e:
                        print(e)
                nueva_linea=''
                for parte in guardar_cadena:
                    nueva_linea=nueva_linea+parte
                #
                try:
                    doc.paragraphs[i].text=''
                    #del doc.paragraphs[i]
                    #doc.paragraphs.pop(i).text
                    #par= doc.add_paragraph()
                    doc.paragraphs[i]=doc.add_paragraph()
                    #run= par.add_run(nueva_linea)
                    run=doc.paragraphs[i].add_run(nueva_linea)
                    font = run.font
                    font.name = 'Arial Narrow'
                    font.size = Pt(12)
                    #del doc.paragraphs[i]

                except Exception as e:
                    print(e)
                #
                #doc.paragraphs[i].text=nueva_linea 
    porcentaje_tablas(doc) 

    doc.save(ruta_guardar) 

def porcentaje_decimal(doc):#version obsoleta
    l=cantidad_lineas(doc)
    a=-1
    lst=[]
    for i in range(l):
        cadena=doc.paragraphs[i].text
        guardar_cadena=cadena.split('%')
        for pos,char in enumerate(cadena):
            if(char == '%'):
                lst.append(pos)
                #a=a+1
        #r=cadena.find('%')#posicion en que se encuentra lo que se busca, en caso de no encontrarse devuelve -1
        longitud=len(lst)
        if (longitud>0):
            for r in lst:
            #if r!=a:
                numero=''
                contador=1
                cond=True
                while(cond):
                    valor=cadena[r-contador]#obtenner el valor
                    cond=valor.isdigit()#comprobar si es digito o no, si es falso acaba la ejecucion
                    if (cond):
                        numero=valor+numero#guardando el numero
                        contador=contador+1#avanza para el siguiente valor
                    elif(valor=='.'):# para casos → ab.cd%
                        contador=contador+1
                        numero='.'+numero
                        cond=True
                    elif (valor==' ' and numero==''):# para casos → ab.cd %
                        contador=contador+1
                        cond=True
                try:
                    punto=numero.find('.')
                    if punto!=a:
                        numero_separado=numero.split('.')#separa el string segun el caracter separador
                        numeroconv1=int(numero_separado[0])#parte entera
                        numeroconv2=int(numero_separado[1])#parte decimal
                        porcentaje1=numero_to_letras(numeroconv1)
                        porcentaje2=numero_to_letras(numeroconv2)
                        porcentajeminus1=porcentaje1.lower()
                        porcentajeminus2=porcentaje2.lower()
                        nuevostring='%('+porcentajeminus1+' coma '+porcentajeminus2+' por ciento)'
                        nueva_cadena=cadena.replace('%',nuevostring,1)
                        doc.paragraphs[i].text=nueva_cadena
                    else:
                        numeroconv1=int(numero)
                        porcentaje1=numero_to_letras(numeroconv1)
                        porcentajeminus1=porcentaje1.lower()
                        nuevostring='%('+porcentajeminus1+' por ciento)'
                        nueva_cadena=cadena.replace('%',nuevostring)
                        doc.paragraphs[i].text=nueva_cadena
                except Exception as e:
                    print(e)
        lst.clear() #vaciar la lista  

    doc.save('C:/Users/DELL/Desktop/pruebaword/pruebadecimalesrepetidos.docx')          

def porcentaje_general(doc,ruta_guardar):#version obsoleta
    l=cantidad_lineas(doc)
    a=-1
    contador=1
    cond=True
    numero1=''
    numero2=''
    m=1
    n=1
    for i in range(l):
        cadena=doc.paragraphs[i].text
        r=cadena.find('%')#posicion en que se encuentra la coma
        if r!=a:
            while(cond):
                num=cadena[r-contador]
                cond=num.isdigit()
                contador=contador+1
                m=m+1
                numero1=num+numero1
            cond=True
            contador=1
            while(cond):
                num=cadena[r-m-contador-1]
                cond=num.isdigit()
                contador=contador+1
                n=n+1
                numero2=num+numero2
        numeroconv1=int(numero1)    
        numeroconv2=int(numero2)
        porcentaje1=numero_to_letras(numeroconv1)
        porcentajeminus1=porcentaje1.lower()
        porcentaje2=numero_to_letras(numeroconv2)
        porcentajeminus2=porcentaje2.lower()
        nuevostring='%('+porcentajeminus2+' punto '+porcentajeminus1+' por ciento)'
        print(nuevostring)
    doc.save(ruta_guardar)        

def editar_linea(doc,string,linea_cambiar,ruta_guardar):
    l=cantidad_lineas(doc)
    a=-1
    c=0
    for i in range(l):
        cadena=doc.paragraphs[i].text
        r=cadena.find(string)
        if r!=a:
            nlinea=cadena.replace(string,linea_cambiar)
            doc.paragraphs[i].text=nlinea
            c=c+1
    if (c==0):
        print('No se encontro')
    else:
        doc.save(ruta_guardar)

def evitar_sobreescritura(cadena,string):#ponerlo dentro de un if
    cond=True
    a=-1
    r=cadena.find(string)#string = 'por ciento'
    if r!=a:
        cond=False
    return cond

def principal(expediente,documento,nombre,string):#por corregir
    #error, str no tiene atributo paragraphs
    doc=f'C:/Users/DELL/Desktop/pruebaword/{expediente}/{documento}/{nombre}.docx'
    #docu=str(doc)
    try:
        buscar_palabra(string,doc)
        cambiarcomas(doc)
        cambiar_porcentaje(doc)
        camrbiarmillon(doc)
    except Exception as e:
        print(e)
    doc.save(f'C:/Users/DELL/Desktop/pruebaword/{expediente}/{documento}/{nombre}modificado.docx')

def principal2(doc,string):#por corregir
    #error, str no tiene atributo paragraphs
    buscar_palabra(string,doc)
    cambiarcomas(doc)
    cambiar_porcentaje(doc)
    camrbiarmillon(doc)
    doc.save('C:/Users/DELL/Desktop/pruebaword/modificado.docx')

def cantidad_lineas(doc):
    l=len(doc.paragraphs)
    return l

def cambiar_porcentaje(doc,ruta_guardar):#version obsoleta
    l=cantidad_lineas(doc)
    a=-1
    for i in range(l):
        cadena=doc.paragraphs[i].text
        r=cadena.find('%')#posicion en que se encuentra la coma
        if r!=a:
            num1=cadena[r-1]
            cond1=num1.isdigit()
            numero=num1
            if (cond1):
                num2=cadena[r-2]
                cond2=num2.isdigit()
                if(cond2):
                    numero=num2+num1
                    num3=cadena[r-3]
                    cond3=num3.isdigit()
                    if (cond3):
                        numero=num3+num2+num1
                    """elif(cadena[r-3]=='.'):
                        num4=cadena[r-4] 
                        cond4=num4.isdigit()  
                        if (cond4):
                            nummero2=num4
                            num5=cadena[r-5]
                            cond5=num5.isdigit()
                            if (cond5):
                                numero2=num5+num4"""
            #print(numero)     
            try:      
                numeroconv=int(numero)
                #intnum=int(numero)
                porcentaje=numero_to_letras(numeroconv)
                porcentajeminus=porcentaje.lower()
                """if (cond4):
                    numeroconv2=int(numero2)
                    porcentaje2=numero_to_letras(numeroconv2)
                    porcentajeminus2=porcentaje2.lower()
                    nuevostring='%('+porcentajeminus2+' punto '+porcentajeminus+' por ciento)'"""
                nuevostring='%('+porcentajeminus+' por ciento)'
                nueva_cadena=cadena.replace('%',nuevostring)
                doc.paragraphs[i].text=nueva_cadena
            except Exception as e:
                print(e)
    porcentaje_tablas(doc)
    doc.save(ruta_guardar)

def buscar_palabra(string,doc):#debe devolver un booleano
    l=cantidad_lineas(doc)
    a=-1
    c=0
    valor=False
    string=string.upper()
    string = reconstruir_string(string)
    for i in range(l):
        cadena=doc.paragraphs[i].text
        r=cadena.find(string)
        if r!=a:
            c=c+1
            valor=True
            break
    if (c==0):
        valor=buscar_en_tabla(string,doc)
    return valor

def cambiarcomas(doc,ruta_guardar):
    l=cantidad_lineas(doc)
    a=-1
    caracteres=['S/','US$']
    for i in range(l):
        cadena=doc.paragraphs[i].text
        for caracter in caracteres:
            r=cadena.find(caracter)#posicion en que se encuentra la coma
            if r!=a:
                nueva_cadena=cadena.replace(',','.')
                doc.paragraphs[i].text=nueva_cadena
    comas_en_tabla(doc)
    doc.save(ruta_guardar)

def camrbiarmillon(doc,ruta_guardar):
    l=cantidad_lineas(doc)
    a=-1
    caracteres=['S/','US$']
    for i in range(l):
        cadena=doc.paragraphs[i].text
        for caracter in caracteres:
            r=cadena.find(caracter)#posicion en que se encuentra la coma
            if r!=a:
                #print('Se encontro en la linea '+str(i+1))
                #print(doc.paragraphs[i].text)
                nueva_cadena=cadena.replace('´','.')
                #print(nueva_cadena)
                doc.paragraphs[i].text=nueva_cadena
    comas_en_tabla(doc)
    doc.save(ruta_guardar)

def buscar_en_tabla(string,doc):
    leertabla=doc.tables
    resultado=0
    a=-1
    cond=False
    for x in range(len(leertabla)):
        tabla=leertabla[x]#obtener la primera tabla
        for i in range(0,len(tabla.rows)):#filas
            for j in range(0,len(tabla.columns)):#columnas
                #print(tabla.cell(i,j).text)
                cadena=tabla.cell(i,j).text
                r=cadena.find(string)
                if (r!=a):
                    print('Se encontro')
                    resultado=resultado+1
                    cond=True
                    break
    if (resultado==0):
        return cond

def comas_en_tabla(doc):
    leertabla=doc.tables
    a=-1
    caracteres=['S/','US$']
    for x in range(len(leertabla)):
        tabla=leertabla[x]#obtener la primera tabla
        for i in range(0,len(tabla.rows)):#filas
            for j in range(0,len(tabla.columns)):#columnas
                #print(tabla.cell(i,j).text)
                cadena=tabla.cell(i,j).text
                #nc=cadena[20:30]
                #print(nc)
                for caracter in caracteres:
                    r=cadena.find(caracter)
                    if (r!=a):
                        nueva_cadena=cadena.replace(',','.')
                        #print(nueva_cadena)
                        tabla.cell(i,j).text=nueva_cadena

def porcentaje_tablas(doc):
    leertabla=doc.tables
    a=-1
    for x in range(len(leertabla)):
        tabla=leertabla[x]#obtener la primera tabla
        for i in range(0,len(tabla.rows)):#filas
            for j in range(0,len(tabla.columns)):#columnas
                #print(tabla.cell(i,j).text)
                cadena=tabla.cell(i,j).text
                #nc=cadena[20:30]
                #print(nc)
                if (evitar_sobreescritura(cadena,'por ciento')):
                    r=cadena.find('%')
                    if (r!=a):
                #print('Se encontro en la linea '+str(i+1))
                #print(doc.paragraphs[i].text)
                        num1=cadena[r-1]
                        cond1=num1.isdigit()
                        numero=num1
                        if (cond1):
                            num2=cadena[r-2]
                            cond2=num2.isdigit()
                            if(cond2):
                                numero=num2+num1
                                num3=cadena[r-3]
                                cond3=num3.isdigit()
                                if (cond3):
                                    numero=num3+num2+num1
                        #print(numero)   
                        try:      
                            numeroconv=int(numero)
                            #intnum=int(numero)
                            porcentaje=numero_to_letras(numeroconv)
                            #porcentaje=numero_to_letras(numero)
                            porcentajeminus=porcentaje.lower()
                            nuevostring='%('+porcentajeminus+' por ciento)'
                            nueva_cadena=cadena.replace('%',nuevostring)
                            tabla.cell(i,j).text=nueva_cadena
                        except Exception as e:
                            print(e)

def agregar_parrafo(string,doc):
    p=doc.add_paragraph(string)
    return p

def pdf_a_word(pdf):
    pdfn=pdf.replace('.pdf','')
    doc = aw.Document(pdf)
    doc.save(f'{pdfn}.docx')

def leerpdf(pdf,palabra):
    pdf_documento = pdf
    documento=fitz.open(pdf_documento)
    p=documento.page_count
    a=-1
    c=0
    for i in range(p):
        pagina=documento.load_page(i)
        text=pagina.get_text('text')
        r=text.find(palabra)
        if r!=a:
            #print('Se encontro en la linea '+str(i+1))
            #print('Se encontro')
            c=c+1
    if (c==0):
        print('No se encontro')
    else:
        print('Se encontro '+str(c)+ ' veces')

def extraerTexto_imagen(palabra,imagen):
    image_path =imagen
    path_to_tesseract = r'C:/Program Files (x86)/Tesseract-OCR/tesseract.exe'
    img = Image.open(image_path) 
    pytesseract.tesseract_cmd = path_to_tesseract 
    text = pytesseract.image_to_string(img) 
    r=text.find(palabra) 
    if r!=-1:
        print('Se encontro')
    else:
        print('No encontro')

def pdf_imagen(pdf):
    # import module
    pages = convert_from_path(pdf)
    for i in range(len(pages)):
        pages[i].save('page'+ str(i) +'.jpg', 'JPEG')

def extraerTexto_imagenV2(nombre,cantidad,palabra,tipodoc,expediente):#modificar la ruta de la imagen
    for i in range(cantidad):
        imagen='C:/Users/DELL/Desktop/pruebaword/'+expediente+'/'+tipodoc+'/'+nombre+f'{i+1}.jpg'
        print('Pagina '+str(i+1)) 
        extraerTexto_imagen(palabra,imagen)

